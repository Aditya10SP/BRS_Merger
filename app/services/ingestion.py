"""
Document ingestion service for parsing PDF and DOCX files.
Converts raw documents into structured BRS and CR models.
"""
import re
from pathlib import Path
from typing import List, Tuple, Optional
from datetime import datetime
import uuid

from pypdf import PdfReader
from docx import Document as DocxDocument

from app.models.schemas import (
    BRSDocument, BRSSection, DocumentMetadata, SectionMetadata,
    ChangeRequest, ChangeDelta, DocumentType, ChangeType,
    ApprovalStatus, Priority
)
from app.core.logging_config import logger


class DocumentParser:
    """Parses PDF and DOCX documents into structured models."""
    
    # Regex patterns for section detection
    SECTION_PATTERN = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)$')
    HEADER_PATTERNS = [
        re.compile(r'^#+\s+(.+)$'),  # Markdown headers
        re.compile(r'^([A-Z][A-Z\s]+)$'),  # ALL CAPS headers
    ]
    
    def __init__(self):
        self.current_doc_id = None
        self.current_version = None
    
    def parse_pdf(self, file_path: Path, doc_type: DocumentType) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            doc_type: Type of document (BRS or CR)
        
        Returns:
            Extracted text content
        """
        logger.info(f"Parsing PDF: {file_path}")
        
        try:
            reader = PdfReader(str(file_path))
            text_content = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    text_content.append(text)
                    logger.debug(f"Extracted {len(text)} chars from page {page_num}")
            
            full_text = "\n".join(text_content)
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
    
    def parse_docx(self, file_path: Path, doc_type: DocumentType) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            doc_type: Type of document (BRS or CR)
        
        Returns:
            Extracted text content
        """
        logger.info(f"Parsing DOCX: {file_path}")
        
        try:
            doc = DocxDocument(str(file_path))
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            full_text = "\n".join(text_content)
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: Path, doc_type: DocumentType) -> str:
        """
        Extract text from PDF or DOCX file.
        
        Args:
            file_path: Path to document
            doc_type: Type of document
        
        Returns:
            Extracted text
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return self.parse_pdf(file_path, doc_type)
        elif suffix in ['.docx', '.doc']:
            return self.parse_docx(file_path, doc_type)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def detect_sections(self, text: str) -> List[Tuple[str, str, str]]:
        """
        Detect sections in text using heuristics.
        
        Args:
            text: Full document text
        
        Returns:
            List of (section_path, title, content) tuples
        """
        sections = []
        lines = text.split('\n')
        
        current_section_path = None
        current_title = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for numbered sections (e.g., "1.2.3 Section Title")
            section_match = self.SECTION_PATTERN.match(line)
            
            if section_match:
                # Save previous section
                if current_section_path and current_title:
                    sections.append((
                        current_section_path,
                        current_title,
                        '\n'.join(current_content).strip()
                    ))
                
                # Start new section
                current_section_path = section_match.group(1)
                current_title = section_match.group(2).strip()
                current_content = []
                logger.debug(f"Detected section: {current_section_path} - {current_title}")
            else:
                # Add to current section content
                if current_section_path:
                    current_content.append(line)
        
        # Save last section
        if current_section_path and current_title:
            sections.append((
                current_section_path,
                current_title,
                '\n'.join(current_content).strip()
            ))
        
        logger.info(f"Detected {len(sections)} sections")
        return sections
    
    def parse_brs_document(
        self,
        file_path: Path,
        doc_id: Optional[str] = None,
        version: Optional[str] = None
    ) -> BRSDocument:
        """
        Parse a BRS document into structured format.
        
        Args:
            file_path: Path to BRS file
            doc_id: Optional document ID (auto-generated if not provided)
            version: Optional version string (auto-detected if not provided)
        
        Returns:
            Structured BRSDocument
        """
        logger.info(f"Parsing BRS document: {file_path}")
        
        # Extract text
        text = self.extract_text(file_path, DocumentType.BRS)
        
        # Generate metadata
        if not doc_id:
            doc_id = f"BRS-{uuid.uuid4().hex[:8]}"
        if not version:
            version = self._detect_version(text) or "v1.0"
        
        metadata = DocumentMetadata(
            doc_id=doc_id,
            doc_type=DocumentType.BRS,
            version=version,
            source_file=file_path.name
        )
        
        # Detect and parse sections
        section_data = self.detect_sections(text)
        sections = []
        
        for section_path, title, content in section_data:
            section_id = f"{doc_id}-SEC-{section_path.replace('.', '-')}"
            
            section_metadata = SectionMetadata(
                section_id=section_id,
                section_title=title,
                section_path=section_path,
                parent_section_id=None  # TODO: Implement hierarchy detection
            )
            
            section = BRSSection(
                metadata=section_metadata,
                content=content,
                subsections=[]
            )
            sections.append(section)
        
        brs_doc = BRSDocument(
            metadata=metadata,
            sections=sections
        )
        
        logger.info(f"Successfully parsed BRS: {doc_id} with {len(sections)} sections")
        return brs_doc
    
    def parse_change_request(
        self,
        file_path: Path,
        cr_id: Optional[str] = None,
        priority: Priority = Priority.MEDIUM,
        approval_status: ApprovalStatus = ApprovalStatus.PENDING
    ) -> ChangeRequest:
        """
        Parse a Change Request document.
        
        Args:
            file_path: Path to CR file
            cr_id: Optional CR ID (auto-generated if not provided)
            priority: Priority level
            approval_status: Approval status
        
        Returns:
            Structured ChangeRequest
        """
        logger.info(f"Parsing Change Request: {file_path}")
        
        # Extract text
        text = self.extract_text(file_path, DocumentType.CHANGE_REQUEST)
        
        # Generate CR ID if not provided
        if not cr_id:
            cr_id = f"CR-{uuid.uuid4().hex[:8]}"
        
        # Extract title (first non-empty line or from pattern)
        title = self._extract_cr_title(text) or f"Change Request {cr_id}"
        
        # Parse deltas (this is a simplified version - in production, this would be more sophisticated)
        deltas = self._parse_cr_deltas(text, cr_id)
        
        cr = ChangeRequest(
            cr_id=cr_id,
            title=title,
            priority=priority,
            approval_status=approval_status,
            source_file=file_path.name,
            deltas=deltas
        )
        
        logger.info(f"Successfully parsed CR: {cr_id} with {len(deltas)} deltas")
        return cr
    
    def _detect_version(self, text: str) -> Optional[str]:
        """Attempt to detect version string from document text."""
        version_patterns = [
            re.compile(r'version\s*[:=]?\s*v?(\d+\.\d+(?:\.\d+)?)', re.IGNORECASE),
            re.compile(r'v(\d+\.\d+(?:\.\d+)?)', re.IGNORECASE),
        ]
        
        for pattern in version_patterns:
            match = pattern.search(text[:1000])  # Search in first 1000 chars
            if match:
                return f"v{match.group(1)}"
        
        return None
    
    def _extract_cr_title(self, text: str) -> Optional[str]:
        """Extract CR title from text."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            # Return first substantial line
            for line in lines[:5]:
                if len(line) > 10 and not line.startswith('#'):
                    return line
        return None
    
    def _parse_cr_deltas(self, text: str, cr_id: str) -> List[ChangeDelta]:
        """
        Parse change deltas from CR text.
        This is a simplified heuristic-based parser.
        In production, this would use more sophisticated NLP or structured templates.
        """
        deltas = []
        
        # Look for patterns like "Section X.Y: Change description"
        section_change_pattern = re.compile(
            r'section\s+(\d+(?:\.\d+)*)[:\s]+(.+?)(?=section\s+\d+|$)',
            re.IGNORECASE | re.DOTALL
        )
        
        matches = section_change_pattern.findall(text)
        
        for idx, (section_path, change_desc) in enumerate(matches):
            delta_id = f"{cr_id}-DELTA-{idx+1:03d}"
            
            # Heuristic: detect change type from keywords
            change_type = ChangeType.MODIFY
            if any(kw in change_desc.lower() for kw in ['add', 'new', 'introduce']):
                change_type = ChangeType.ADD
            elif any(kw in change_desc.lower() for kw in ['remove', 'delete', 'eliminate']):
                change_type = ChangeType.DELETE
            
            delta = ChangeDelta(
                delta_id=delta_id,
                impacted_section_id=f"SEC-{section_path.replace('.', '-')}",
                impacted_section_title=f"Section {section_path}",
                change_type=change_type,
                old_content=None,  # Would need more sophisticated parsing
                new_content=change_desc.strip(),
                rationale="Extracted from CR document"
            )
            deltas.append(delta)
        
        # If no structured changes found, create a single generic delta
        if not deltas:
            delta = ChangeDelta(
                delta_id=f"{cr_id}-DELTA-001",
                impacted_section_id="UNKNOWN",
                impacted_section_title="General Changes",
                change_type=ChangeType.MODIFY,
                old_content=None,
                new_content=text[:500],  # First 500 chars as summary
                rationale="Full CR content - requires manual review"
            )
            deltas.append(delta)
        
        return deltas
