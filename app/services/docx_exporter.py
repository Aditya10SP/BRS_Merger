"""
DOCX export service for generating formatted BRS documents.
"""
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.schemas import FinalBRS, GeneratedSection
from app.core.logging_config import logger


class DOCXExporter:
    """Export Final BRS to formatted DOCX document."""
    
    def export_to_docx(self, final_brs: FinalBRS, output_path: Path) -> None:
        """
        Export Final BRS to DOCX document.
        
        Args:
            final_brs: Final BRS object
            output_path: Path to save DOCX file
        """
        logger.info(f"Exporting BRS to DOCX: {output_path}")
        
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = final_brs.title or "Business Requirements Specification"
        doc.core_properties.author = "BRS Consolidation System"
        doc.core_properties.subject = f"BRS {final_brs.version}"
        
        # Add title page
        self._create_title_page(doc, final_brs)
        doc.add_page_break()
        
        # Add metadata section
        self._create_metadata_section(doc, final_brs)
        
        # Add sections
        for section in final_brs.sections:
            self._create_section(doc, section)
        
        # Add validation notes if present
        if final_brs.validation_notes:
            doc.add_page_break()
            self._create_validation_section(doc, final_brs)
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info(f"DOCX exported successfully: {output_path}")
    
    def _create_title_page(self, doc: Document, final_brs: FinalBRS) -> None:
        """Create title page."""
        # Add spacing
        for _ in range(5):
            doc.add_paragraph()
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(final_brs.title or "Business Requirements Specification")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(26, 26, 26)
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Scramble 2.0 - Sensitive Data Masking System")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(26, 26, 26)
        
        doc.add_paragraph()
        
        # Metadata table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = "Document ID:"
        table.cell(0, 1).text = final_brs.brs_id
        table.cell(1, 0).text = "Version:"
        table.cell(1, 1).text = final_brs.version
        table.cell(2, 0).text = "Generated:"
        table.cell(2, 1).text = datetime.now().strftime('%B %d, %Y')
        table.cell(3, 0).text = "Status:"
        table.cell(3, 1).text = "Final"
        
        # Make first column bold
        for i in range(4):
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
    
    def _create_metadata_section(self, doc: Document, final_brs: FinalBRS) -> None:
        """Create metadata section."""
        # Source documents
        if final_brs.source_brs_documents:
            p = doc.add_paragraph()
            run = p.add_run("Source BRS Documents:")
            run.font.bold = True
            run.font.size = Pt(12)
            
            for doc_id in final_brs.source_brs_documents:
                doc.add_paragraph(f"• {doc_id}", style='List Bullet')
        
        # Applied CRs
        if final_brs.applied_change_requests:
            p = doc.add_paragraph()
            run = p.add_run("Applied Change Requests:")
            run.font.bold = True
            run.font.size = Pt(12)
            
            for cr_id in final_brs.applied_change_requests:
                doc.add_paragraph(f"• {cr_id}", style='List Bullet')
        
        # Validation status
        p = doc.add_paragraph()
        run = p.add_run("Validation Status: ")
        run.font.bold = True
        run.font.size = Pt(12)
        
        status_run = p.add_run("✓ Passed" if final_brs.validation_passed else "✗ Failed")
        status_run.font.color.rgb = RGBColor(0, 128, 0) if final_brs.validation_passed else RGBColor(255, 0, 0)
        status_run.font.bold = True
        
        doc.add_paragraph()
    
    def _create_section(self, doc: Document, section: GeneratedSection) -> None:
        """Create section."""
        # Section heading
        heading_text = f"{section.section_path} {section.section_title}"
        level = section.section_path.count('.') + 1
        
        if level == 1:
            heading = doc.add_heading(heading_text, level=1)
        elif level == 2:
            heading = doc.add_heading(heading_text, level=2)
        else:
            heading = doc.add_heading(heading_text, level=3)
        
        # Section content
        paragraphs = section.content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.paragraph_format.space_after = Pt(12)
        
        # Add traceability information
        if section.source_documents or section.applied_changes:
            trace_parts = []
            if section.source_documents:
                trace_parts.append(f"Sources: {', '.join(section.source_documents)}")
            if section.applied_changes:
                trace_parts.append(f"Applied Changes: {', '.join(section.applied_changes)}")
            
            trace_text = " | ".join(trace_parts)
            p = doc.add_paragraph(f"[{trace_text}]")
            run = p.runs[0]
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(102, 102, 102)
    
    def _create_validation_section(self, doc: Document, final_brs: FinalBRS) -> None:
        """Create validation notes section."""
        doc.add_heading("Validation Notes", level=1)
        
        for note in final_brs.validation_notes[:20]:
            doc.add_paragraph(f"• {note}", style='List Bullet')
        
        if len(final_brs.validation_notes) > 20:
            p = doc.add_paragraph(f"... and {len(final_brs.validation_notes) - 20} more notes")
            p.runs[0].font.italic = True
